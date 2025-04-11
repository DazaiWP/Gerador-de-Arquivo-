import os
import ttkbootstrap as ttkb
from ttkbootstrap.constants import *
from tkinter import messagebox
from docx import Document
from datetime import datetime
import subprocess
import win32com.client

# Função segura para Desktop
def get_desktop_path():
    try:
        from winreg import OpenKey, QueryValueEx, HKEY_CURRENT_USER
        key = OpenKey(HKEY_CURRENT_USER, r"Software\Microsoft\Windows\CurrentVersion\Explorer\Shell Folders")
        return QueryValueEx(key, "Desktop")[0]
    except Exception:
        return os.path.join(os.path.expanduser("~"), "Desktop")

# Caminho dos modelos e log
PASTA_MODELOS = r"S:\TOTVS\09 - TERMO DE RESPONSABILIDADE\GERADOR\MODELO"
LOG_PATH = r"S:\TOTVS\09 - TERMO DE RESPONSABILIDADE\GERADOR\Logs\termos_gerados.log"

modelos_disponiveis = [f for f in os.listdir(PASTA_MODELOS) if f.endswith(".docx")]
if not modelos_disponiveis:
    raise FileNotFoundError("Nenhum modelo .docx encontrado na pasta.")

caminho_arquivo_gerado = None

# Funções de formatação
def formatar_data(data_str):
    data_str = data_str.strip()
    if data_str.isdigit() and len(data_str) == 8:
        return f"{data_str[:2]}/{data_str[2:4]}/{data_str[4:]}"
    return data_str

def formatar_cpf(cpf_str):
    cpf_str = cpf_str.strip()
    if cpf_str.isdigit() and len(cpf_str) == 11:
        return f"{cpf_str[:3]}.{cpf_str[3:6]}.{cpf_str[6:9]}-{cpf_str[9:]}"
    return cpf_str

def formatar_telefone(telefone_str):
    telefone_str = telefone_str.strip()
    if telefone_str.isdigit():
        if len(telefone_str) == 11:
            return f"({telefone_str[:2]}) {telefone_str[2:7]}-{telefone_str[7:]}"
        elif len(telefone_str) == 10:
            return f"({telefone_str[:2]}) {telefone_str[2:6]}-{telefone_str[6:]}"
    return telefone_str

def formatar_cartao(cartao_str):
    cartao_str = cartao_str.strip()
    if cartao_str.isdigit() and len(cartao_str) == 16:
        return f"{cartao_str[:4]} {cartao_str[4:8]} {cartao_str[8:12]} {cartao_str[12:]}"
    return cartao_str

# Substituição de texto nos documentos
def substituir_texto(doc, substituicoes):
    def processa_texto(texto, subs):
        for chave, valor in subs.items():
            if isinstance(valor, list):
                for v in valor:
                    if chave in texto:
                        texto = texto.replace(chave, v, 1)
            else:
                texto = texto.replace(chave, valor)
        return texto

    for p in doc.paragraphs:
        novo_texto = processa_texto(p.text, substituicoes)
        if novo_texto != p.text:
            p.clear()
            p.add_run(novo_texto)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    novo_texto = processa_texto(paragrafo.text, substituicoes)
                    if novo_texto != paragrafo.text:
                        paragrafo.clear()
                        paragrafo.add_run(novo_texto)

# Enviar por e-mail
def enviar_email():
    if not caminho_arquivo_gerado:
        messagebox.showwarning("Aviso", "Nenhum termo gerado ainda.")
        return
    destinatario = entry_email.get().strip()
    if not destinatario:
        messagebox.showerror("Erro", "Preencha o e-mail do destinatário.")
        return
    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        mail = outlook.CreateItem(0)
        mail.To = destinatario
        mail.Subject = "Termo de Responsabilidade"
        mail.Body = (
            "Prezados (a),\n\n"
            "Segue anexo o termo de responsabilidade referente ao equipamento entregue.\n\n"
            "Atenciosamente,"
        )
        mail.Attachments.Add(caminho_arquivo_gerado)
        mail.Display()
    except Exception as e:
        messagebox.showerror("Erro ao enviar e-mail", str(e))

# Geração do termo
def gerar_documento():
    global caminho_arquivo_gerado
    modelo_selecionado = combo_modelo.get()
    if not modelo_selecionado:
        messagebox.showerror("Erro", "Selecione um modelo.")
        return

    nome = entry_nome.get().strip()
    cpf = formatar_cpf(entry_cpf.get())
    funcao = entry_funcao.get().strip()
    data = formatar_data(entry_data.get())
    marca = entry_marca.get().strip()
    modelo_valor = entry_modelo.get().strip()
    serie = entry_serie.get().strip()
    telefone = formatar_telefone(entry_telefone.get())
    cartao = formatar_cartao(entry_cartao.get())

    caminho_modelo = os.path.join(PASTA_MODELOS, modelo_selecionado)
    try:
        doc = Document(caminho_modelo)
    except Exception as e:
        messagebox.showerror("Erro", f"Erro ao abrir o modelo:\n{e}")
        return

    dados = {
        "[ nome do colaborador ]": nome,
        "[ CPF do colaborador ]": cpf,
        "[ função do colaborador ]": funcao,
        "[ data do dia]": data,
        "[ marca do notebook ]": marca,
        "[ modelo do notebook ]": modelo_valor,
        "[ Nº de Serie do notebook ]": serie,
        "[ marca do celular ]": marca,
        "[ modelo do celular ]": modelo_valor,
        "[ número de série ]": serie,
        "[ número de telefone ]": telefone,
        "[ número do cartão ]": cartao,
        "[informar]": [marca, modelo_valor, serie, telefone],
    }

    substituir_texto(doc, dados)

    nome_arquivo = f"Termo_{nome.replace(' ', '_')}.docx"
    caminho = os.path.join(get_desktop_path(), nome_arquivo)

    try:
        doc.save(caminho)
        caminho_arquivo_gerado = caminho
    except Exception as e:
        messagebox.showerror("Erro ao salvar", str(e))
        return

    try:
        usuario = os.getlogin()
    except:
        usuario = "desconhecido"

    os.makedirs(os.path.dirname(LOG_PATH), exist_ok=True)
    with open(LOG_PATH, "a", encoding="utf-8") as log:
        log.write(f"[{datetime.now().strftime('%d/%m/%Y %H:%M')}] Quem gerou: {usuario} | Modelo: {modelo_selecionado} | Nome: {nome} | CPF: {cpf} | Arquivo: {caminho}\n")

    messagebox.showinfo("Sucesso", f"✅ Termo gerado como:\n{nome_arquivo}")
    botao_abrir.config(state=NORMAL)

# Interface gráfica
root = ttkb.Window(themename="flatly")
root.title("Gerador de Termo de Responsabilidade - Luma Service")
root.geometry("700x550")
root.resizable(True, True)

mainframe = ttkb.Frame(root, padding=20)
mainframe.pack(fill="both", expand=True)

ttkb.Label(mainframe, text="Modelo:", font=("Segoe UI", 11)).grid(row=0, column=0, sticky="W", pady=5)
combo_modelo = ttkb.Combobox(mainframe, values=modelos_disponiveis, state="readonly", width=50)
combo_modelo.grid(row=0, column=1, columnspan=2, pady=5, sticky="EW")

campos = [
    "Nome do colaborador*",
    "CPF*",
    "Função*",
    "Data*",
    "Marca (Celular/Notebook):",
    "Modelo:",
    "Número de Série:",
    "Telefone:",
    "Número do Cartão:",
    "E-mail do destinatário:",
]

entries = []
for i, campo in enumerate(campos, start=1):
    ttkb.Label(mainframe, text=campo, font=("Segoe UI", 10)).grid(row=i, column=0, sticky="W", pady=3)
    entry = ttkb.Entry(mainframe, width=50)
    entry.grid(row=i, column=1, columnspan=2, sticky="EW", pady=3)
    entries.append(entry)

(
    entry_nome, entry_cpf, entry_funcao, entry_data,
    entry_marca, entry_modelo, entry_serie,
    entry_telefone, entry_cartao, entry_email
) = entries

botao_gerar = ttkb.Button(mainframe, text="Gerar Termo", command=gerar_documento, bootstyle=PRIMARY, width=20)
botao_gerar.grid(row=len(campos)+1, column=1, pady=10, sticky="EW")

botao_abrir = ttkb.Button(mainframe, text="Abrir Documento", command=lambda: os.startfile(caminho_arquivo_gerado) if caminho_arquivo_gerado else None, bootstyle=INFO, state=DISABLED, width=20)
botao_abrir.grid(row=len(campos)+2, column=1, pady=5, sticky="EW")

botao_email = ttkb.Button(mainframe, text="Enviar por E-mail", command=enviar_email, bootstyle=SUCCESS, width=20)
botao_email.grid(row=len(campos)+3, column=1, pady=5, sticky="EW")

mainframe.columnconfigure(1, weight=1)
mainframe.columnconfigure(2, weight=1)

root.mainloop()
